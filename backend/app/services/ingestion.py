import os
from typing import Dict
import pypdf
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import pdf2image
import io

async def process_uploaded_file(file_path: str, original_filename: str, file_extension: str) -> Dict:
    """Process uploaded file and extract content"""
    
    result = {
        "title": original_filename,
        "content": "",
        "summary": "",
        "metadata": {}
    }
    
    try:
        if file_extension == "pdf":
            result = await _process_pdf(file_path, original_filename)
        elif file_extension == "docx":
            result = await _process_docx(file_path, original_filename)
        elif file_extension == "txt":
            result = await _process_txt(file_path, original_filename)
        elif file_extension in ["png", "jpg", "jpeg"]:
            result = await _process_image(file_path, original_filename)
        else:
            result["content"] = "Binary file - text extraction not supported"
    
    except Exception as e:
        print(f"Error processing file: {e}")
        result["content"] = f"Error processing file: {str(e)}"
    
    return result

async def _process_pdf(file_path: str, original_filename: str) -> Dict:
    """Extract text from PDF file with OCR fallback for scanned PDFs"""
    content = ""
    ocr_used = False
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            # Extract text from all pages
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        
        # Check if extracted text is too short (likely scanned PDF)
        if len(content.strip()) < 50 and len(pdf_reader.pages) > 0:
            print("PDF appears to be scanned, attempting OCR...")
            content = await _ocr_pdf(file_path)
            ocr_used = True
        
        # Try to extract metadata
        try:
            pdf_info = pdf_reader.metadata
            title = pdf_info.get('/Title', original_filename)
        except:
            title = original_filename
        
        metadata = {
            "pages": len(pdf_reader.pages),
            "format": "PDF",
            "ocr_used": ocr_used
        }
        
        return {
            "title": title,
            "content": content.strip(),
            "summary": content[:500] if content else "",
            "metadata": metadata
        }
    
    except Exception as e:
        print(f"Error processing PDF: {e}")
        # Try OCR as fallback
        try:
            print("Attempting OCR as fallback...")
            content = await _ocr_pdf(file_path)
            return {
                "title": original_filename,
                "content": content.strip(),
                "summary": content[:500] if content else "",
                "metadata": {
                    "format": "PDF",
                    "ocr_used": True,
                    "extraction_method": "ocr_fallback"
                }
            }
        except Exception as ocr_error:
            print(f"OCR fallback also failed: {ocr_error}")
            return {
                "title": original_filename,
                "content": f"Error processing PDF: {str(e)}",
                "summary": "",
                "metadata": {}
            }

async def _ocr_pdf(file_path: str) -> str:
    """Perform OCR on PDF by converting pages to images"""
    try:
        # Convert PDF to images
        images = pdf2image.convert_from_path(file_path)
        
        content = ""
        for i, image in enumerate(images):
            # Perform OCR on each page
            text = pytesseract.image_to_string(image)
            content += f"\n--- Page {i+1} ---\n{text}\n"
        
        return content.strip()
    except Exception as e:
        print(f"Error during PDF OCR: {e}")
        raise e

async def _process_docx(file_path: str, original_filename: str) -> Dict:
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Try to get title from first paragraph
        title = original_filename
        if doc.paragraphs:
            first_paragraph = doc.paragraphs[0].text.strip()
            if first_paragraph:
                title = first_paragraph
        
        return {
            "title": title,
            "content": content.strip(),
            "summary": content[:500] if content else "",
            "metadata": {
                "paragraphs": len(doc.paragraphs),
                "format": "DOCX"
            }
        }
    
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return {
            "title": original_filename,
            "content": f"Error processing DOCX: {str(e)}",
            "summary": "",
            "metadata": {}
        }

async def _process_txt(file_path: str, original_filename: str) -> Dict:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Use filename as title, or first line if available
        lines = content.split('\n')
        title = original_filename
        if lines and lines[0].strip():
            title = lines[0].strip()
        
        return {
            "title": title,
            "content": content.strip(),
            "summary": content[:500] if content else "",
            "metadata": {
                "lines": len(lines),
                "format": "TXT"
            }
        }
    
    except Exception as e:
        print(f"Error processing TXT: {e}")
        return {
            "title": original_filename,
            "content": f"Error processing TXT: {str(e)}",
            "summary": "",
            "metadata": {}
        }

async def _process_image(file_path: str, original_filename: str) -> Dict:
    """Extract text from image using OCR"""
    try:
        # Open image
        image = Image.open(file_path)
        
        # Perform OCR
        content = pytesseract.image_to_string(image)
        
        return {
            "title": original_filename,
            "content": content.strip(),
            "summary": content[:500] if content else "",
            "metadata": {
                "format": "Image",
                "size": image.size
            }
        }
    
    except Exception as e:
        print(f"Error processing image: {e}")
        return {
            "title": original_filename,
            "content": f"Error processing image (OCR failed): {str(e)}",
            "summary": "",
            "metadata": {}
        }